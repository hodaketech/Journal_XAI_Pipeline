from __future__ import annotations

import argparse

from common import OUT_DIR, ensure_dirs, read_csv, require_file
from docx.shared import Pt


def add_caption(document, text: str) -> None:
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True


def add_table(document, rows: list[dict[str, str]], columns: list[str]) -> None:
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, column in enumerate(columns):
        header[idx].text = column
    for row in rows:
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            cells[idx].text = str(row.get(column, ""))
    for table_row in table.rows:
        for cell in table_row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8)
    document.add_paragraph()


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Appendix B Word tables from generated CSV files.")
    parser.add_argument("--out", default=str(OUT_DIR / "appendix_b_tables.docx"))
    args = parser.parse_args()
    ensure_dirs()
    require_file(OUT_DIR / "table_B1.csv", "Run survey/scripts/03_extract_report.py first.")
    require_file(OUT_DIR / "table_B2.csv", "Run survey/scripts/03_extract_report.py first.")

    from docx import Document
    from docx.enum.section import WD_ORIENT
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(9)

    table_b1 = read_csv(OUT_DIR / "table_B1.csv")
    table_b2 = read_csv(OUT_DIR / "table_B2.csv")
    add_caption(document, "Table B.1: Record-level extraction for retained records.")
    add_table(document, table_b1, list(table_b1[0].keys()) if table_b1 else [])
    add_caption(document, "Table B.2: Derived distributions over the retained corpus.")
    add_table(document, table_b2, ["group", "category", "count"])
    document.save(args.out)
    print(f"Wrote {args.out}")


if __name__ == "__main__":
    main()
